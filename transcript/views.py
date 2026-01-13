# import os
# import re
# import whisper
# from django.http import JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from django.conf import settings
# from pymongo import MongoClient

# # ================= MongoDB =================
# client = MongoClient("mongodb://localhost:27017/")
# db = client["audio_transcript_db"]
# collection = db["audio_results"]

# # ================= Whisper =================
# model = whisper.load_model("base")


# # ================= Helper =================
# def normalize(word: str) -> str:
#     """lowercase + remove punctuation"""
#     return re.sub(r"[^\w]", "", word.lower())


# # ================= Main API =================
# @csrf_exempt
# def process_audio(request):
#     if request.method != "POST":
#         return JsonResponse({"error": "POST required"}, status=400)

#     audio_file = request.FILES.get("audio")
#     template_name = request.POST.get("template")

#     if not audio_file or not template_name:
#         return JsonResponse(
#             {"error": "audio or template missing"},
#             status=400
#         )

#     # -------- STEP 1: SAVE AUDIO --------
#     audio_dir = os.path.join(settings.MEDIA_ROOT, "audio")
#     os.makedirs(audio_dir, exist_ok=True)

#     audio_path = os.path.join(audio_dir, audio_file.name)
#     with open(audio_path, "wb+") as f:
#         for chunk in audio_file.chunks():
#             f.write(chunk)

#     # -------- STEP 2: TRANSCRIBE --------
#     result = model.transcribe(
#         audio_path,
#         word_timestamps=True
#     )

#     language = result.get("language")

#     # -------- STEP 3: BUILD FULL TRANSCRIPT ARRAY --------
#     transcript = []
#     for seg in result["segments"]:
#         for w in seg.get("words", []):
#             transcript.append({
#                 "word": normalize(w["word"]),
#                 "start": round(w["start"], 3),
#                 "end": round(w["end"], 3)
#             })

#     # -------- STEP 4: READ TEMPLATE --------
#     template_path = os.path.join(
#         settings.MEDIA_ROOT,
#         "templates_text",
#         template_name
#     )

#     if not os.path.exists(template_path):
#         return JsonResponse(
#             {"error": "Template file not found"},
#             status=404
#         )

#     with open(template_path, "r") as f:
#         template_words = {
#             normalize(w)
#             for w in f.read().split()
#             if normalize(w)
#         }

#     # -------- STEP 5: FIND COMMON WORDS --------
#     common_words = []

#     for w in transcript:
#         if w["word"] in template_words:
#             common_words.append({
#                 "word": w["word"],
#                 "start": w["start"],
#                 "end": w["end"]
#             })

#     # -------- STEP 6: SAVE TO MONGODB --------
#     doc = {
#         "audio": audio_file.name,
#         "language": language,
#         "transcript": transcript,
#         "common_words": common_words
#     }

#     inserted = collection.insert_one(doc)

#     # -------- STEP 7: RESPONSE --------
#     return JsonResponse(
#         {
#             "_id": str(inserted.inserted_id),
#             "audio": audio_file.name,
#             "language": language,
#             "transcript": transcript,
#             "common_words": common_words
#         },
#         json_dumps_params={"indent": 2}
#     )

# import os
# import re
# import whisper
# from datetime import datetime
# from django.http import JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from django.conf import settings
# from pymongo import MongoClient
# from pydub import AudioSegment

# # ================= MongoDB =================
# client = MongoClient("mongodb://localhost:27017/")
# db = client["audio_transcript_db"]
# collection = db["audio_results"]

# # ================= Whisper =================
# model = whisper.load_model("base")


# # ================= Helper =================
# def normalize(word: str) -> str:
#     return re.sub(r"[^\w]", "", word.lower())


# # ================= Main API =================
# @csrf_exempt
# def process_audio(request):
#     if request.method != "POST":
#         return JsonResponse({"error": "POST required"}, status=400)

#     audio_file = request.FILES.get("audio")
#     template_name = request.POST.get("template")

#     if not audio_file or not template_name:
#         return JsonResponse(
#             {"error": "audio or template missing"},
#             status=400
#         )

#     # -------- STEP 1: SAVE AUDIO --------
#     audio_dir = os.path.join(settings.MEDIA_ROOT, "audio")
#     os.makedirs(audio_dir, exist_ok=True)

#     audio_path = os.path.join(audio_dir, audio_file.name)
#     with open(audio_path, "wb+") as f:
#         for chunk in audio_file.chunks():
#             f.write(chunk)

#     # -------- STEP 2: TRANSCRIBE --------
#     result = model.transcribe(audio_path, word_timestamps=True)
#     language = result.get("language")

#     # -------- STEP 3: FULL TRANSCRIPT --------
#     transcript = []
#     for seg in result["segments"]:
#         for w in seg.get("words", []):
#             transcript.append({
#                 "word": normalize(w["word"]),
#                 "start": round(w["start"], 3),
#                 "end": round(w["end"], 3)
#             })

#     # -------- STEP 4: READ TEMPLATE --------
#     template_path = os.path.join(
#         settings.MEDIA_ROOT,
#         "templates_text",
#         template_name
#     )

#     if not os.path.exists(template_path):
#         return JsonResponse(
#             {"error": "Template file not found"},
#             status=404
#         )

#     with open(template_path, "r") as f:
#         template_words = {
#             normalize(w)
#             for w in f.read().split()
#             if normalize(w)
#         }

#     # -------- STEP 5: COMMON WORDS --------
#     common_words = [
#         w for w in transcript if w["word"] in template_words
#     ]

#     # -------- STEP 6: CREATE DB RECORD FIRST --------
#     base_name = os.path.splitext(audio_file.name)[0]

#     insert_result = collection.insert_one({
#         "audio": audio_file.name,
#         "template": template_name,
#         "language": language,
#         "created_at": datetime.utcnow(),
#         "transcript": transcript,
#         "common_words": common_words,
#         "files": {}
#     })

#     record_id = str(insert_result.inserted_id)

#     # -------- STEP 7: SAVE UNIQUE TEXT FILE --------
#     common_dir = os.path.join(settings.MEDIA_ROOT, "common_words")
#     os.makedirs(common_dir, exist_ok=True)

#     common_txt_filename = f"{base_name}_{record_id}_common_words.txt"
#     common_txt_path = os.path.join(common_dir, common_txt_filename)

#     with open(common_txt_path, "w") as f:
#         for w in common_words:
#             f.write(f"{w['word']} -> {w['start']}s - {w['end']}s\n")

#     # -------- STEP 8: BUILD UNIQUE COMBINED AUDIO --------
#     original_audio = AudioSegment.from_file(audio_path)
#     combined_audio = AudioSegment.silent(duration=0)

#     for w in common_words:
#         combined_audio += original_audio[
#             int(w["start"] * 1000):int(w["end"] * 1000)
#         ]

#     clips_dir = os.path.join(settings.MEDIA_ROOT, "audio_clips")
#     os.makedirs(clips_dir, exist_ok=True)

#     combined_audio_filename = f"{base_name}_{record_id}_common_words.wav"
#     combined_audio_path = os.path.join(
#         clips_dir, combined_audio_filename
#     )

#     combined_audio.export(combined_audio_path, format="wav")

#     # -------- STEP 9: UPDATE DB WITH FILE NAMES --------
#     collection.update_one(
#         {"_id": insert_result.inserted_id},
#         {"$set": {
#             "files.common_audio": combined_audio_filename,
#             "files.common_text": common_txt_filename
#         }}
#     )

#     # -------- STEP 10: RESPONSE --------
#     return JsonResponse(
#         {
#             "_id": record_id,
#             "audio": audio_file.name,
#             "template": template_name,
#             "language": language,
#             "transcript": transcript,
#             "common_words": common_words,
#             "common_audio": combined_audio_filename,
#             "common_text": common_txt_filename
#         },
#         json_dumps_params={"indent": 2}
#     )

# import os
# import re
# from datetime import datetime

# import whisper
# from django.http import JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from django.conf import settings
# from pymongo import MongoClient
# from pydub import AudioSegment


# # ================= MongoDB =================
# client = MongoClient("mongodb://localhost:27017/")
# db = client["audio_transcript_db"]
# collection = db["audio_results"]


# # ================= Whisper =================
# model = whisper.load_model("base")


# # ================= Helpers =================
# def normalize(word: str) -> str:
#     return re.sub(r"[^\w]", "", word.lower())


# # ================= Main API =================
# @csrf_exempt
# def process_audio(request):
#     if request.method != "POST":
#         return JsonResponse({"error": "POST required"}, status=400)

#     audio_file = request.FILES.get("audio")
#     template_name = request.POST.get("template")

#     if not audio_file or not template_name:
#         return JsonResponse(
#             {"error": "audio or template missing"},
#             status=400
#         )

#     # -------- STEP 1: SAVE AUDIO --------
#     audio_dir = os.path.join(settings.MEDIA_ROOT, "audio")
#     os.makedirs(audio_dir, exist_ok=True)

#     audio_path = os.path.join(audio_dir, audio_file.name)
#     with open(audio_path, "wb+") as f:
#         for chunk in audio_file.chunks():
#             f.write(chunk)

#     # -------- STEP 2: TRANSCRIBE --------
#     result = model.transcribe(audio_path, word_timestamps=True)
#     language = result.get("language")

#     # -------- STEP 3: FULL WORD TRANSCRIPT --------
#     transcript = []
#     for seg in result["segments"]:
#         for w in seg.get("words", []):
#             transcript.append({
#                 "word": normalize(w["word"]),
#                 "start": round(w["start"], 3),
#                 "end": round(w["end"], 3)
#             })

#     # -------- STEP 4: LOAD TEMPLATE --------
#     template_path = os.path.join(
#         settings.MEDIA_ROOT,
#         "templates_text",
#         template_name
#     )

#     if not os.path.exists(template_path):
#         return JsonResponse(
#             {"error": "Template file not found"},
#             status=404
#         )

#     with open(template_path, "r") as f:
#         template_words = [
#             normalize(w)
#             for w in f.read().split()
#             if normalize(w)
#         ]

#     # -------- STEP 5: MATCH WORDS (TEMPLATE ORDER) --------
#     matched_words = []

#     for t_word in template_words:
#         for w in transcript:
#             if w["word"] == t_word:
#                 matched_words.append({
#                     "word": t_word,
#                     "start": w["start"],
#                     "end": w["end"]
#                 })
#                 break

#     # -------- STEP 6: CREATE DB RECORD --------
#     base_name = os.path.splitext(audio_file.name)[0]

#     insert_result = collection.insert_one({
#         "audio": audio_file.name,
#         "template": template_name,
#         "language": language,
#         "created_at": datetime.utcnow(),
#         "transcript": transcript,
#         "matched_words": [],
#         "phrases": []
#     })

#     record_id = str(insert_result.inserted_id)

#     # -------- STEP 7: WORD-LEVEL AUDIO CLIPS --------
#     clips_dir = os.path.join(settings.MEDIA_ROOT, "audio_clips")
#     os.makedirs(clips_dir, exist_ok=True)

#     original_audio = AudioSegment.from_file(audio_path)
#     word_clips = []

#     for idx, w in enumerate(matched_words):
#         clip = original_audio[
#             int(w["start"] * 1000):int(w["end"] * 1000)
#         ]

#         clip_name = f"{base_name}_{record_id}_{w['word']}_{idx}.wav"
#         clip_path = os.path.join(clips_dir, clip_name)

#         clip.export(clip_path, format="wav")

#         word_clips.append({
#             "word": w["word"],
#             "start": w["start"],
#             "end": w["end"],
#             "audio_file": clip_name
#         })

#     # -------- STEP 8: GROUP INTO PHRASES --------
#     phrases_raw = []
#     current = []

#     for i in range(len(word_clips)):
#         if i == 0:
#             current.append(word_clips[i])
#             continue

#         prev = word_clips[i - 1]
#         curr = word_clips[i]

#         if curr["start"] - prev["end"] < 0.5:
#             current.append(curr)
#         else:
#             phrases_raw.append(current)
#             current = [curr]

#     if current:
#         phrases_raw.append(current)

#     # -------- STEP 9: PHRASE-LEVEL AUDIO --------
#     phrase_results = []

#     for idx, phrase in enumerate(phrases_raw):
#         start = phrase[0]["start"]
#         end = phrase[-1]["end"]

#         clip = original_audio[
#             int(start * 1000):int(end * 1000)
#         ]

#         phrase_text = " ".join(p["word"] for p in phrase)
#         phrase_name = f"{base_name}_{record_id}_phrase_{idx}.wav"

#         clip.export(
#             os.path.join(clips_dir, phrase_name),
#             format="wav"
#         )

#         phrase_results.append({
#             "texttext": phrase_text,
#             "start": start,
#             "end": end,
#             "audio_file": phrase_name
#         })

#     # -------- STEP 10: UPDATE DB --------
#     collection.update_one(
#         {"_id": insert_result.inserted_id},
#         {"$set": {
#             "matched_words": word_clips,
#             "phrases": phrase_results
#         }}
#     )

#     # -------- STEP 11: RESPONSE --------
#     return JsonResponse(
#         {
#             "_id": record_id,
#             "audio": audio_file.name,
#             "template": template_name,
#             "language": language,
#             "matched_words": word_clips,
#             "phrases": phrase_results
#         },
#         json_dumps_params={"indent": 2}
#     )


# correct version:
# import os
# import re
# from datetime import datetime

# import whisper
# from django.http import JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from django.conf import settings
# from pymongo import MongoClient
# from pydub import AudioSegment


# # ================= MongoDB =================
# client = MongoClient("mongodb://localhost:27017/")
# db = client["audio_transcript_db"]
# collection = db["audio_results"]


# # ================= Whisper =================
# model = whisper.load_model("base")


# # ================= Helpers =================
# def normalize(word: str) -> str:
#     # Lowercase and remove punctuation
#     return re.sub(r"[^\w]", "", word.lower())


# # ================= Main API =================
# @csrf_exempt
# def process_audio(request):
#     if request.method != "POST":
#         return JsonResponse({"error": "POST required"}, status=400)

#     audio_file = request.FILES.get("audio")
#     template_name = request.POST.get("template")

#     if not audio_file or not template_name:
#         return JsonResponse(
#             {"error": "audio or template missing"},
#             status=400
#         )

#     # -------- STEP 1: SAVE AUDIO --------
#     audio_dir = os.path.join(settings.MEDIA_ROOT, "audio")
#     os.makedirs(audio_dir, exist_ok=True)

#     audio_path = os.path.join(audio_dir, audio_file.name)
#     with open(audio_path, "wb+") as f:
#         for chunk in audio_file.chunks():
#             f.write(chunk)

#     # -------- STEP 2: TRANSCRIBE --------
#     # word_timestamps=True is crucial for segmentation
#     result = model.transcribe(audio_path, word_timestamps=True)
#     language = result.get("language")

#     # -------- STEP 3: FULL WORD TRANSCRIPT --------
#     transcript = []
#     # We flatten the segments into a single list of words with their original index
#     global_index = 0 
#     for seg in result["segments"]:
#         for w in seg.get("words", []):
#             transcript.append({
#                 "index": global_index,  # ID to track continuity
#                 "word": normalize(w["word"]),
#                 "start": round(w["start"], 3),
#                 "end": round(w["end"], 3)
#             })
#             global_index += 1

#     # -------- STEP 4: LOAD TEMPLATE --------
#     template_path = os.path.join(
#         settings.MEDIA_ROOT,
#         "templates_text",
#         template_name
#     )

#     if not os.path.exists(template_path):
#         return JsonResponse(
#             {"error": "Template file not found"},
#             status=404
#         )

#     with open(template_path, "r") as f:
#         # Normalize template words
#         template_words = [
#             normalize(w)
#             for w in f.read().split()
#             if normalize(w)
#         ]

#     # -------- STEP 5: MATCH WORDS (STRICT SEQUENTIAL) --------
#     # Logic: Find template words in the transcript in the correct order.
    
#     matched_words = []
#     transcript_ptr = 0
    
#     for t_word in template_words:
#         # Search for this template word in the remaining transcript
#         for i in range(transcript_ptr, len(transcript)):
#             if transcript[i]["word"] == t_word:
#                 # Found a match
#                 matched_words.append(transcript[i])
                
#                 # Move pointer to the next word for the next iteration
#                 transcript_ptr = i + 1 
#                 break
        
#         # If the transcript runs out, stop matching
#         if transcript_ptr >= len(transcript):
#             break

#     # -------- STEP 6: CREATE DB RECORD (INITIAL) --------
#     base_name = os.path.splitext(audio_file.name)[0]

#     insert_result = collection.insert_one({
#         "audio": audio_file.name,
#         "template": template_name,
#         "language": language,
#         "created_at": datetime.utcnow(),
#         "transcript": transcript,
#         "matched_words": [],
#         "phrases": []
#     })

#     record_id = str(insert_result.inserted_id)

#     # -------- STEP 7: EXPORT WORD-BY-WORD AUDIO --------
#     clips_dir = os.path.join(settings.MEDIA_ROOT, "audio_clips")
#     os.makedirs(clips_dir, exist_ok=True)

#     original_audio = AudioSegment.from_file(audio_path)
    
#     word_results = []
    
#     # We loop through matched_words to export individual clips
#     for idx, w in enumerate(matched_words):
#         clip = original_audio[
#             int(w["start"] * 1000):int(w["end"] * 1000)
#         ]

#         clip_name = f"{base_name}_{record_id}_word_{idx}_{w['word']}.wav"
#         clip_path = os.path.join(clips_dir, clip_name)
#         clip.export(clip_path, format="wav")

#         word_results.append({
#             "word": w["word"],
#             "start": w["start"],
#             "end": w["end"],
#             "audio_file": clip_name
#         })

#     # -------- STEP 8: GROUPING LOGIC (MISMATCH & 20s LIMIT) --------
#     # Requirements:
#     # 1. Contiguous matching (no words spoken in between).
#     # 2. Duration < 20 seconds.
#     # 3. If mismatch (skipped word or extra word), break.

#     phrases_raw = []
#     if matched_words:
#         current_group = [matched_words[0]]

#         for i in range(1, len(matched_words)):
#             prev = current_group[-1]
#             curr = matched_words[i]

#             # Check 1: Continuity (Mismatch Detection)
#             # The indices in the original transcript must be sequential (e.g., 5 and 6).
#             # If they are 5 and 7, the speaker said an extra word at 6, so we split.
#             is_contiguous = (curr["index"] == prev["index"] + 1)

#             # Check 2: Duration Limit
#             # Start of group to end of current word
#             group_duration = curr["end"] - current_group[0]["start"]
#             is_within_limit = group_duration < 20.0

#             if is_contiguous and is_within_limit:
#                 current_group.append(curr)
#             else:
#                 # Close current group and start new
#                 phrases_raw.append(current_group)
#                 current_group = [curr]
        
#         # Append the last remaining group
#         if current_group:
#             phrases_raw.append(current_group)

#     # -------- STEP 9: EXPORT PHRASE AUDIO (COMMON AUDIO) --------
#     phrase_results = []

#     for idx, phrase in enumerate(phrases_raw):
#         start_time = phrase[0]["start"]
#         end_time = phrase[-1]["end"]
        
#         # Join text for "Common Text"
#         phrase_text = " ".join(p["word"] for p in phrase)

#         # Cut Audio
#         clip = original_audio[
#             int(start_time * 1000):int(end_time * 1000)
#         ]

#         phrase_name = f"{base_name}_{record_id}_phrase_{idx}.wav"
#         phrase_path = os.path.join(clips_dir, phrase_name)
        
#         clip.export(phrase_path, format="wav")

#         phrase_results.append({
#             "common_text": phrase_text,
#             "start": start_time,
#             "end": end_time,
#             "duration": round(end_time - start_time, 2),
#             "common_audio_file": phrase_name
#         })

#     # -------- STEP 10: UPDATE DB --------
#     collection.update_one(
#         {"_id": insert_result.inserted_id},
#         {"$set": {
#             "matched_words": word_results,
#             "phrases": phrase_results
#         }}
#     )

#     # -------- STEP 11: RESPONSE --------
#     return JsonResponse(
#         {
#             "_id": record_id,
#             "status": "success",
#             "counts": {
#                 "words": len(word_results),
#                 "phrases": len(phrase_results)
#             },
#             # "matched_words": word_results, # Uncomment if you want full list in response
#             "common_data": phrase_results
#         },
#         json_dumps_params={"indent": 2}
#     )


# import os
# import re
# from datetime import datetime

# import whisper
# from django.http import JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from django.conf import settings
# from pymongo import MongoClient
# from pydub import AudioSegment


# # ================= MongoDB =================
# client = MongoClient("mongodb://localhost:27017/")
# db = client["audio_transcript_db"]
# collection = db["audio_results"]


# # ================= Whisper =================
# model = whisper.load_model("base")


# # ================= Helpers =================
# def normalize(word: str) -> str:
#     return re.sub(r"[^\w]", "", word.lower())

# def format_time_info(start, end, text):
#     """Helper to format content for the text file"""
#     duration = round(end - start, 3)
#     return (
#         f"Start: {start:.3f}\n"
#         f"End: {end:.3f}\n"
#         f"Duration: {duration:.3f}\n"
#         f"Text: {text}"
#     )


# # ================= Main API =================
# @csrf_exempt
# def process_audio(request):
#     if request.method != "POST":
#         return JsonResponse({"error": "POST required"}, status=400)

#     audio_file = request.FILES.get("audio")
#     template_name = request.POST.get("template")

#     if not audio_file or not template_name:
#         return JsonResponse(
#             {"error": "audio or template missing"},
#             status=400
#         )

#     # -------- STEP 1: SAVE AUDIO --------
#     audio_dir = os.path.join(settings.MEDIA_ROOT, "audio")
#     os.makedirs(audio_dir, exist_ok=True)

#     audio_path = os.path.join(audio_dir, audio_file.name)
#     with open(audio_path, "wb+") as f:
#         for chunk in audio_file.chunks():
#             f.write(chunk)

#     # -------- STEP 2: TRANSCRIBE --------
#     result = model.transcribe(audio_path, word_timestamps=True)
#     language = result.get("language")

#     # -------- STEP 3: FULL WORD TRANSCRIPT --------
#     transcript = []
#     global_index = 0
#     for seg in result["segments"]:
#         for w in seg.get("words", []):
#             start = round(w["start"], 3)
#             end = round(w["end"], 3)
#             duration = round(end - start, 3)
            
#             transcript.append({
#                 "index": global_index,
#                 "word": normalize(w["word"]),
#                 "raw_word": w["word"].strip(),
#                 "start": start,
#                 "end": end,
#                 "duration": duration  # Added duration here
#             })
#             global_index += 1

#     # -------- STEP 4: LOAD TEMPLATE --------
#     template_path = os.path.join(
#         settings.MEDIA_ROOT,
#         "templates_text",
#         template_name
#     )

#     if not os.path.exists(template_path):
#         return JsonResponse({"error": "Template file not found"}, status=404)

#     with open(template_path, "r") as f:
#         template_words = [
#             normalize(w) for w in f.read().split() if normalize(w)
#         ]

#     # -------- STEP 5: MATCH WORDS (STRICT SEQUENTIAL) --------
#     matched_words_raw = []
#     transcript_ptr = 0
    
#     for t_word in template_words:
#         for i in range(transcript_ptr, len(transcript)):
#             if transcript[i]["word"] == t_word:
#                 matched_words_raw.append(transcript[i])
#                 transcript_ptr = i + 1
#                 break
#         if transcript_ptr >= len(transcript):
#             break

#     # -------- STEP 6: CREATE DB RECORD --------
#     base_name = os.path.splitext(audio_file.name)[0]

#     insert_result = collection.insert_one({
#         "audio": audio_file.name,
#         "template": template_name,
#         "language": language,
#         "created_at": datetime.utcnow(),
#         "transcript": transcript,
#         "matched": [] 
#     })

#     record_id = str(insert_result.inserted_id)

#     # -------- STEP 7: EXPORT WORD FILES (Audio & Detail Text) --------
#     clips_dir = os.path.join(settings.MEDIA_ROOT, "audio_clips")
#     os.makedirs(clips_dir, exist_ok=True)

#     original_audio = AudioSegment.from_file(audio_path)
    
#     for idx, w in enumerate(matched_words_raw):
#         clip = original_audio[int(w["start"] * 1000):int(w["end"] * 1000)]
#         clip_name = f"{base_name}_{record_id}_word_{idx}_{w['word']}.wav"
#         txt_name = clip_name.replace(".wav", ".txt")
        
#         # 1. Export Audio
#         clip.export(os.path.join(clips_dir, clip_name), format="wav")
        
#         # 2. Export Text with Timestamps
#         content = format_time_info(w["start"], w["end"], w["raw_word"])
#         with open(os.path.join(clips_dir, txt_name), "w", encoding="utf-8") as tf:
#             tf.write(content)

#     # -------- STEP 8: GROUPING LOGIC (20s Limit & Continuity) --------
#     grouped_phrases = []
#     if matched_words_raw:
#         current_group = [matched_words_raw[0]]

#         for i in range(1, len(matched_words_raw)):
#             prev = current_group[-1]
#             curr = matched_words_raw[i]

#             is_contiguous = (curr["index"] == prev["index"] + 1)
#             group_duration = curr["end"] - current_group[0]["start"]
#             is_within_limit = group_duration < 20.0

#             if is_contiguous and is_within_limit:
#                 current_group.append(curr)
#             else:
#                 grouped_phrases.append(current_group)
#                 current_group = [curr]
        
#         if current_group:
#             grouped_phrases.append(current_group)

#     # -------- STEP 9: EXPORT MATCHED SEGMENTS (Audio & Detail Text) --------
#     matched_results = []

#     for idx, group in enumerate(grouped_phrases):
#         start_time = group[0]["start"]
#         end_time = group[-1]["end"]
#         duration = round(end_time - start_time, 3)
#         text_content = " ".join(p["raw_word"] for p in group)

#         # 1. Cut Audio
#         clip = original_audio[int(start_time * 1000):int(end_time * 1000)]
        
#         audio_filename = f"{base_name}_{record_id}_matched_{idx}.wav"
#         text_filename = audio_filename.replace(".wav", ".txt")
        
#         # 2. Export Audio
#         clip.export(os.path.join(clips_dir, audio_filename), format="wav")
        
#         # 3. Export Text with Timestamps
#         file_content = format_time_info(start_time, end_time, text_content)
#         with open(os.path.join(clips_dir, text_filename), "w", encoding="utf-8") as tf:
#             tf.write(file_content)

#         # JSON Object
#         matched_results.append({
#             "text": text_content,
#             "start": start_time,
#             "end": end_time,
#             "duration": duration,
#             "audio_file": audio_filename,
#             "text_file": text_filename
#         })

#     # -------- STEP 10: UPDATE DB --------
#     collection.update_one(
#         {"_id": insert_result.inserted_id},
#         {"$set": {
#             "matched": matched_results
#         }}
#     )

#     # -------- STEP 11: RESPONSE --------
#     return JsonResponse(
#         {
#             "_id": record_id,
#             "transcript": transcript,
#             "matched": matched_results
#         },
#         json_dumps_params={"indent": 2}
#     )


# last version is correct.


# import os
# import re

# #difflib use panna smart-aa match aagum

# import difflib  # <--- NEW IMPORT for smart matching
# from datetime import datetime

# import whisper
# from django.http import JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from django.conf import settings
# from pymongo import MongoClient
# from pydub import AudioSegment


# # ================= MongoDB =================
# client = MongoClient("mongodb://localhost:27017/")
# db = client["audio_transcript_db"]
# collection = db["audio_results"]


# # ================= Whisper =================
# model = whisper.load_model("base")


# # ================= Helpers =================
# def normalize(word: str) -> str:
#     # Lowercase and remove punctuation for better matching
#     return re.sub(r"[^\w]", "", word.lower())

# def format_time_info(start, end, text):
#     """Helper to format content for the text file"""
#     duration = round(end - start, 3)
#     return (
#         f"Start: {start:.3f}\n"
#         f"End: {end:.3f}\n"
#         f"Duration: {duration:.3f}\n"
#         f"Text: {text}"  
#     )


# # ================= Main API =================
# @csrf_exempt
# def process_audio(request):
#     if request.method != "POST":
#         return JsonResponse({"error": "POST required"}, status=400)

#     audio_file = request.FILES.get("audio")
#     template_name = request.POST.get("template")

#     if not audio_file or not template_name:
#         return JsonResponse(
#             {"error": "audio or template missing"},
#             status=400
#         )

#     # -------- STEP 1: SAVE AUDIO --------
#     audio_dir = os.path.join(settings.MEDIA_ROOT, "audio")
#     os.makedirs(audio_dir, exist_ok=True)

#     audio_path = os.path.join(audio_dir, audio_file.name)
#     with open(audio_path, "wb+") as f:
#         for chunk in audio_file.chunks():
#             f.write(chunk)

#     # -------- STEP 2: TRANSCRIBE --------
#     result = model.transcribe(audio_path, word_timestamps=True)
#     language = result.get("language")
# #----------------------break-----------------------------------
#     # -------- STEP 3: PREPARE TRANSCRIPT --------
#     transcript = []
#     global_index = 0
    
#     # We create a simple list of words for the matcher to use later
#     transcript_words_list = [] 

#     for seg in result["segments"]:
#         for w in seg.get("words", []):
#             start = round(w["start"], 3)
#             end = round(w["end"], 3)
#             duration = round(end - start, 3)
#             norm_word = normalize(w["word"])
            
#             # Store full object
#             obj = {
#                 "index": global_index,
#                 "word": norm_word,
#                 "raw_word": w["word"].strip(),
#                 "start": start,
#                 "end": end,
#                 "duration": duration
#             }
#             transcript.append(obj)
            
#             # Store just the normalized string for comparison
#             transcript_words_list.append(norm_word)
            
#             global_index += 1

#     # -------- STEP 4: LOAD TEMPLATE --------
#     template_path = os.path.join(
#         settings.MEDIA_ROOT,
#         "templates_text",
#         template_name
#     )

#     if not os.path.exists(template_path):
#         return JsonResponse({"error": "Template file not found"}, status=404)

#     with open(template_path, "r") as f:
#         # Create list of normalized template words
#         template_words_list = [
#             normalize(w) for w in f.read().split() if normalize(w)
#         ]

#     # -------- STEP 5: SMART MATCHING (DIFFLIB) --------
#     # This replaces the old greedy loop. It finds the best alignment.
    
#     matched_words_raw = []
    
#     # SequenceMatcher compares the two lists of words
#     matcher = difflib.SequenceMatcher(None, template_words_list, transcript_words_list)
    
#     # get_opcodes returns blocks like: ('equal', i1, i2, j1, j2)
#     # i is for template, j is for transcript
#     for tag, i1, i2, j1, j2 in matcher.get_opcodes():
#         if tag == 'equal':
#             # These indices match!
#             # We take the words from the TRANSCRIPT (j1 to j2) because they have timestamps
#             segment_match = transcript[j1:j2]
#             matched_words_raw.extend(segment_match)

#     # -------- STEP 6: CREATE DB RECORD --------
#     base_name = os.path.splitext(audio_file.name)[0]

#     insert_result = collection.insert_one({
#         "audio": audio_file.name,
#         "template": template_name,
#         "language": language,
#         "created_at": datetime.utcnow(),
#         "transcript": transcript,
#         "matched": [] 
#     })

#     record_id = str(insert_result.inserted_id)

#     # -------- STEP 7: EXPORT INDIVIDUAL WORD FILES --------
#     clips_dir = os.path.join(settings.MEDIA_ROOT, "audio_clips")
#     os.makedirs(clips_dir, exist_ok=True)

#     original_audio = AudioSegment.from_file(audio_path)
    
#     for idx, w in enumerate(matched_words_raw):
#         clip = original_audio[int(w["start"] * 1000):int(w["end"] * 1000)]
#         clip_name = f"{base_name}_{record_id}_word_{idx}_{w['word']}.wav"
#         txt_name = clip_name.replace(".wav", ".txt")
        
#         # Export Audio
#         clip.export(os.path.join(clips_dir, clip_name), format="wav")
        
#         # Export Text Details
#         content = format_time_info(w["start"], w["end"], w["raw_word"])
#         with open(os.path.join(clips_dir, txt_name), "w", encoding="utf-8") as tf:
#             tf.write(content)

#     # -------- STEP 8: GROUPING LOGIC --------
#     # Groups the raw matched words into phrases if they are contiguous
#     grouped_phrases = []
#     if matched_words_raw:
#         current_group = [matched_words_raw[0]]

#         for i in range(1, len(matched_words_raw)):
#             prev = current_group[-1]
#             curr = matched_words_raw[i]

#             # 1. Check if words are neighbors in the original transcript
#             is_contiguous = (curr["index"] == prev["index"] + 1)
            
#             # 2. Check 20s limit
#             group_duration = curr["end"] - current_group[0]["start"]
#             is_within_limit = group_duration < 20.0

#             if is_contiguous and is_within_limit:
#                 current_group.append(curr)
#             else:
#                 grouped_phrases.append(current_group)
#                 current_group = [curr]
        
#         if current_group:
#             grouped_phrases.append(current_group)

#     # -------- STEP 9: EXPORT MATCHED SEGMENTS (PHRASES) --------
#     matched_results = []

#     for idx, group in enumerate(grouped_phrases):
#         start_time = group[0]["start"]
#         end_time = group[-1]["end"]
#         duration = round(end_time - start_time, 3)
#         text_content = " ".join(p["raw_word"] for p in group)

#         # Cut Audio
#         clip = original_audio[int(start_time * 1000):int(end_time * 1000)]
        
#         audio_filename = f"{base_name}_{record_id}_matched_{idx}.wav"
#         text_filename = audio_filename.replace(".wav", ".txt")
        
#         # Save Files
#         clip.export(os.path.join(clips_dir, audio_filename), format="wav")
        
#         file_content = format_time_info(start_time, end_time, text_content)
#         with open(os.path.join(clips_dir, text_filename), "w", encoding="utf-8") as tf:
#             tf.write(file_content)

#         # Add to Array
#         matched_results.append({
#             "text": text_content,
#             "start": start_time,
#             "end": end_time,
#             "duration": duration,
#             "audio_file": audio_filename,
#             "text_file": text_filename
#         })

#     # -------- STEP 10: UPDATE DB --------
#     collection.update_one(
#         {"_id": insert_result.inserted_id},
#         {"$set": {
#             "matched": matched_results
#         }}
#     )

#     # -------- STEP 11: RESPONSE --------
#     return JsonResponse(
#         {
#             "_id": record_id,
#             "transcript": transcript,
#             "matched": matched_results
#         },
#         json_dumps_params={"indent": 2}
#     )

#29/12/25

# import os
# import re
# import difflib
# from datetime import datetime
# import whisper
# from django.http import JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from django.conf import settings
# from pymongo import MongoClient
# from pydub import AudioSegment
# from docx import Document   


# client = MongoClient("mongodb://localhost:27017/")
# db = client["audio_transcript_db"]
# collection = db["audio_results"]


# model = whisper.load_model("base")



# def normalize(word: str) -> str:
#     return re.sub(r"[^\w]", "", word.lower())


# def format_time_info(start, end, text):
#     duration = round(end - start, 3)
#     return (
#         f"Start: {start:.3f}\n"
#         f"End: {end:.3f}\n"
#         f"Duration: {duration:.3f}\n"
#         f"Text: {text}"
#     )


# def extract_template_text(template_path, ext):
#     """
#     Extract text from TXT / DOC / DOCX / AUDIO template
#     """
    
#     if ext == ".txt":
#         with open(template_path, "r", encoding="utf-8") as f:
#             return f.read()

    
#     elif ext in [".doc", ".docx"]:
#         doc = Document(template_path)
#         return " ".join(p.text for p in doc.paragraphs)

    
#     elif ext in [".wav", ".mp3", ".m4a"]:
#         result = model.transcribe(template_path)
#         return result["text"]

#     else:
#         return None



# @csrf_exempt
# def process_audio(request):
#     if request.method != "POST":
#         return JsonResponse({"error": "POST required"}, status=400)

#     audio_file = request.FILES.get("audio")
#     template_file = request.FILES.get("template")

#     if not audio_file or not template_file:
#         return JsonResponse({"error": "audio or template missing"}, status=400)

#     # -------- STEP 1: SAVE AUDIO --------
#     audio_dir = os.path.join(settings.MEDIA_ROOT, "audio")
#     os.makedirs(audio_dir, exist_ok=True)

#     audio_path = os.path.join(audio_dir, audio_file.name)
#     with open(audio_path, "wb+") as f:
#         for chunk in audio_file.chunks():
#             f.write(chunk)

#     # -------- STEP 2: TRANSCRIBE AUDIO --------
#     result = model.transcribe(audio_path, word_timestamps=True)
#     language = result.get("language")

#     # -------- STEP 3: PREPARE TRANSCRIPT --------
#     transcript = []
#     global_index = 0
#     transcript_words_list = []

#     for seg in result["segments"]:
#         for w in seg.get("words", []):
#             start = round(w["start"], 3)
#             end = round(w["end"], 3)

#             if end <= start:
#                 end = round(start + 0.01, 3)

#             duration = round(end - start, 3)
#             norm_word = normalize(w["word"])

#             obj = {
#                 "index": global_index,
#                 "word": norm_word,
#                 "raw_word": w["word"].strip(),
#                 "start": start,
#                 "end": end,
#                 "duration": duration
#             }

#             transcript.append(obj)
#             transcript_words_list.append(norm_word)
#             global_index += 1

#     # -------- STEP 4: SAVE TEMPLATE FILE --------
#     template_dir = os.path.join(settings.MEDIA_ROOT, "templates")
#     os.makedirs(template_dir, exist_ok=True)

#     template_path = os.path.join(template_dir, template_file.name)
#     with open(template_path, "wb+") as f:
#         for chunk in template_file.chunks():
#             f.write(chunk)

#     ext = os.path.splitext(template_file.name)[1].lower()

#     template_text = extract_template_text(template_path, ext)
#     if not template_text:
#         return JsonResponse(
#             {"error": "Unsupported template format"},
#             status=400
#         )

#     template_words_list = [
#         normalize(w) for w in template_text.split() if normalize(w)
#     ]


#     matched_words_raw = []
#     matcher = difflib.SequenceMatcher(
#         None,
#         template_words_list,
#         transcript_words_list
#     )

#     for tag, i1, i2, j1, j2 in matcher.get_opcodes():
#         if tag == "equal":
#             matched_words_raw.extend(transcript[j1:j2])

   
#     base_name = os.path.splitext(audio_file.name)[0]

#     insert_result = collection.insert_one({
#         "audio": audio_file.name,
#         "template": template_file.name,
#         "language": language,
#         "created_at": datetime.utcnow(),
#         "transcript": transcript,
#         "matched": []
#     })

#     record_id = str(insert_result.inserted_id)

   
#     clips_dir = os.path.join(settings.MEDIA_ROOT, "audio_clips")
#     os.makedirs(clips_dir, exist_ok=True)

#     original_audio = AudioSegment.from_file(audio_path)

#     for idx, w in enumerate(matched_words_raw):
#         start_ms = int(w["start"] * 1000)
#         end_ms = int(w["end"] * 1000)

#         if end_ms <= start_ms:
#             end_ms = start_ms + 100

#         clip = original_audio[start_ms:end_ms]

#         clip_name = f"{base_name}_{record_id}_word_{idx}_{w['word']}.wav"
#         txt_name = clip_name.replace(".wav", ".txt")

#         clip.export(os.path.join(clips_dir, clip_name), format="wav")

#         content = format_time_info(w["start"], w["end"], w["raw_word"])
#         with open(os.path.join(clips_dir, txt_name), "w", encoding="utf-8") as tf:
#             tf.write(content)

  
#     grouped_phrases = []
#     if matched_words_raw:
#         current_group = [matched_words_raw[0]]

#         for i in range(1, len(matched_words_raw)):
#             prev = current_group[-1]
#             curr = matched_words_raw[i]

#             is_contiguous = curr["index"] == prev["index"] + 1
#             group_duration = curr["end"] - current_group[0]["start"]
#             is_within_limit = group_duration < 20.0

#             if is_contiguous and is_within_limit:
#                 current_group.append(curr)
#             else:
#                 grouped_phrases.append(current_group)
#                 current_group = [curr]

#         grouped_phrases.append(current_group)

 
#     matched_results = []

#     for idx, group in enumerate(grouped_phrases):
#         start_time = group[0]["start"]
#         end_time = group[-1]["end"]
#         duration = round(end_time - start_time, 3)
#         text_content = " ".join(p["raw_word"] for p in group)

#         clip = original_audio[
#             int(start_time * 1000):int(end_time * 1000)
#         ]

#         audio_filename = f"{base_name}_{record_id}_matched_{idx}.wav"
#         text_filename = audio_filename.replace(".wav", ".txt")

#         clip.export(os.path.join(clips_dir, audio_filename), format="wav")

#         file_content = format_time_info(start_time, end_time, text_content)
#         with open(os.path.join(clips_dir, text_filename), "w", encoding="utf-8") as tf:
#             tf.write(file_content)

#         matched_results.append({
#             "text": text_content,
#             "start": start_time,
#             "end": end_time,
#             "duration": duration,
#             "audio_file": audio_filename,
#             "text_file": text_filename
#         })

  
#     collection.update_one(
#         {"_id": insert_result.inserted_id},
#         {"$set": {"matched": matched_results}}
#     )

   
#     return JsonResponse(
#         {
#             "_id": record_id,
#             "transcript": transcript,
#             "matched": matched_results
#         },
#         json_dumps_params={"indent": 2}
#     )


#30/12/25

import os
import re
import difflib
from datetime import datetime
import whisper
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from pymongo import MongoClient
from pydub import AudioSegment
from docx import Document


# ================= MongoDB =================
client = MongoClient("mongodb://localhost:27017/")
db = client["audio_transcript_db"]
collection = db["audio_results"]


# ================= Whisper =================
model = whisper.load_model("base")


# ================= Helpers =================
def normalize(word: str) -> str:
    return re.sub(r"[^\w]", "", word.lower())


def format_time_info(start, end, text):
    duration = round(end - start, 3)
    return (
        f"Start: {start:.3f}\n"
        f"End: {end:.3f}\n"
        f"Duration: {duration:.3f}\n"
        f"Text: {text}"
    )


def extract_template_text(template_path, ext):
    if ext == ".txt":
        with open(template_path, "r", encoding="utf-8") as f:
            return f.read()

    elif ext in [".doc", ".docx"]:
        doc = Document(template_path)
        return " ".join(p.text for p in doc.paragraphs)

    elif ext in [".wav", ".mp3", ".m4a"]:
        result = model.transcribe(template_path)
        return result["text"]

    return None


# ================= Main API =================
@csrf_exempt
def process_audio(request):
    if request.method != "POST":
        return JsonResponse({"error": "POST required"}, status=400)

    audio_file = request.FILES.get("audio")
    template_file = request.FILES.get("template")

    if not audio_file or not template_file:
        return JsonResponse({"error": "audio or template missing"}, status=400)

    # -------- STEP 1: SAVE AUDIO --------
    audio_dir = os.path.join(settings.MEDIA_ROOT, "audio")
    os.makedirs(audio_dir, exist_ok=True)

    audio_path = os.path.join(audio_dir, audio_file.name)
    with open(audio_path, "wb+") as f:
        for chunk in audio_file.chunks():
            f.write(chunk)

    # -------- STEP 2: TRANSCRIBE AUDIO --------
    result = model.transcribe(audio_path, word_timestamps=True)
    language = result.get("language")

    # -------- STEP 3: PREPARE TRANSCRIPT --------
    transcript = []
    transcript_words_list = []
    global_index = 0

    for seg in result["segments"]:
        for w in seg.get("words", []):
            start = round(w["start"], 3)
            end = round(w["end"], 3)

            if end <= start:
                end = round(start + 0.01, 3)

            obj = {
                "index": global_index,
                "word": normalize(w["word"]),
                "raw_word": w["word"].strip(),
                "start": start,
                "end": end,
                "duration": round(end - start, 3),
            }

            transcript.append(obj)
            transcript_words_list.append(obj["word"])
            global_index += 1

    # -------- STEP 4: SAVE TEMPLATE --------
    template_dir = os.path.join(settings.MEDIA_ROOT, "templates")
    os.makedirs(template_dir, exist_ok=True)

    template_path = os.path.join(template_dir, template_file.name)
    with open(template_path, "wb+") as f:
        for chunk in template_file.chunks():
            f.write(chunk)

    ext = os.path.splitext(template_file.name)[1].lower()
    template_text = extract_template_text(template_path, ext)

    if not template_text:
        return JsonResponse({"error": "Unsupported template format"}, status=400)

    template_words_list = [
        normalize(w) for w in template_text.split() if normalize(w)
    ]

    # -------- STEP 5: SMART MATCHING --------
    matched_words_raw = []
    matcher = difflib.SequenceMatcher(
        None, template_words_list, transcript_words_list
    )

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            matched_words_raw.extend(transcript[j1:j2])

    # -------- STEP 6: DB INSERT --------
    base_name = os.path.splitext(audio_file.name)[0]

    insert_result = collection.insert_one({
        "audio": audio_file.name,
        "template": template_file.name,
        "language": language,
        "created_at": datetime.utcnow(),
        "transcript": transcript,
        "matched": []
    })

    record_id = str(insert_result.inserted_id)

    # -------- STEP 7: EXPORT WORD CLIPS (SMART SLOW) --------
    clips_dir = os.path.join(settings.MEDIA_ROOT, "audio_clips")
    os.makedirs(clips_dir, exist_ok=True)

    original_audio = AudioSegment.from_file(audio_path)

    for idx, w in enumerate(matched_words_raw):
        start_ms = int(w["start"] * 1000)
        end_ms = int(w["end"] * 1000)

        if end_ms <= start_ms:
            end_ms = start_ms + 150

        clip = original_audio[start_ms:end_ms]  #---orginal audio cut

        # 🎯 ONLY slow very short words
        if clip.duration_seconds < 0.25:
            clip = clip._spawn(
                clip.raw_data,
                overrides={"frame_rate": int(clip.frame_rate * 0.75)} # frame rate kammi paniruke
            ).set_frame_rate(clip.frame_rate)

        # Safety cap 
        if clip.duration_seconds > 3:
            clip = clip[:3000]

        clip_name = f"{base_name}_{record_id}_word_{idx}_{w['word']}.wav"
        txt_name = clip_name.replace(".wav", ".txt")

        clip.export(os.path.join(clips_dir, clip_name), format="wav")

        content = format_time_info(w["start"], w["end"], w["raw_word"])
        with open(os.path.join(clips_dir, txt_name), "w", encoding="utf-8") as tf:
            tf.write(content)

    # -------- STEP 8: GROUPING --------
    grouped_phrases = []

    if matched_words_raw:
        current = [matched_words_raw[0]]

        for i in range(1, len(matched_words_raw)):
            prev = current[-1]
            curr = matched_words_raw[i]

            if (
                curr["index"] == prev["index"] + 1 and
                curr["end"] - current[0]["start"] < 20
            ):
                current.append(curr)
            else:
                grouped_phrases.append(current)
                current = [curr]

        grouped_phrases.append(current)

    # -------- STEP 9: EXPORT PHRASES --------
    matched_results = []

    for idx, group in enumerate(grouped_phrases):
        start = group[0]["start"]
        end = group[-1]["end"]
        text = " ".join(w["raw_word"] for w in group)

        clip = original_audio[int(start * 1000):int(end * 1000)]

        if clip.duration_seconds > 2:
            clip = clip[:2000]

        audio_file_name = f"{base_name}_{record_id}_matched_{idx}.wav"
        text_file_name = audio_file_name.replace(".wav", ".txt")

        clip.export(os.path.join(clips_dir, audio_file_name), format="wav")

        with open(os.path.join(clips_dir, text_file_name), "w", encoding="utf-8") as f:
            f.write(format_time_info(start, end, text))

        matched_results.append({
            "text": text,
            "start": start,
            "end": end,
            "duration": round(end - start, 3),
            "audio_file": audio_file_name,
            "text_file": text_file_name
        })

    # -------- STEP 10: UPDATE DB --------
    collection.update_one(
        {"_id": insert_result.inserted_id},
        {"$set": {"matched": matched_results}}
    )

    # -------- STEP 11: RESPONSE --------
    return JsonResponse(
        {
            "_id": record_id,
            "transcript": transcript,
            "matched": matched_results
        },
        json_dumps_params={"indent": 2}
    )


# manual method "and" " to "


# import os
# import re
# import difflib
# from datetime import datetime
# import whisper
# from django.http import JsonResponse
# from django.views.decorators.csrf import csrf_exempt
# from django.conf import settings
# from pymongo import MongoClient
# from pydub import AudioSegment
# from docx import Document


# # ================= MongoDB =================
# client = MongoClient("mongodb://localhost:27017/")
# db = client["audio_transcript_db"]
# collection = db["audio_results"]


# # ================= Whisper =================
# model = whisper.load_model("base")


# # ================= Helpers =================
# def normalize(word: str) -> str:
#     return re.sub(r"[^\w]", "", word.lower())


# def format_time_info(start, end, text):
#     duration = round(end - start, 3)
#     return (
#         f"Start: {start:.3f}\n"
#         f"End: {end:.3f}\n"
#         f"Duration: {duration:.3f}\n"
#         f"Text: {text}"
#     )


# def extract_template_text(template_path, ext):
#     if ext == ".txt":
#         with open(template_path, "r", encoding="utf-8") as f:
#             return f.read()

#     elif ext in [".doc", ".docx"]:
#         doc = Document(template_path)
#         return " ".join(p.text for p in doc.paragraphs)

#     elif ext in [".wav", ".mp3", ".m4a"]:
#         result = model.transcribe(template_path)
#         return result["text"]

#     return None


# # ================= Main API =================
# @csrf_exempt
# def process_audio(request):
#     if request.method != "POST":
#         return JsonResponse({"error": "POST required"}, status=400)

#     audio_file = request.FILES.get("audio")
#     template_file = request.FILES.get("template")

#     if not audio_file or not template_file:
#         return JsonResponse({"error": "audio or template missing"}, status=400)

#     # -------- STEP 1: SAVE AUDIO --------
#     audio_dir = os.path.join(settings.MEDIA_ROOT, "audio")
#     os.makedirs(audio_dir, exist_ok=True)

#     audio_path = os.path.join(audio_dir, audio_file.name)
#     with open(audio_path, "wb+") as f:
#         for chunk in audio_file.chunks():
#             f.write(chunk)

#     # -------- STEP 2: TRANSCRIBE AUDIO --------
#     result = model.transcribe(audio_path, word_timestamps=True)
#     language = result.get("language")

#     # -------- STEP 3: PREPARE TRANSCRIPT --------
#     transcript = []
#     transcript_words_list = []
#     global_index = 0

#     for seg in result["segments"]:
#         for w in seg.get("words", []):
#             start = round(w["start"], 3)
#             end = round(w["end"], 3)

#             if end <= start:
#                 end = round(start + 0.01, 3)

#             obj = {
#                 "index": global_index,
#                 "word": normalize(w["word"]),
#                 "raw_word": w["word"].strip(),
#                 "start": start,
#                 "end": end,
#                 "duration": round(end - start, 3),
#             }

#             transcript.append(obj)
#             transcript_words_list.append(obj["word"])
#             global_index += 1

#     # -------- STEP 4: SAVE TEMPLATE --------
#     template_dir = os.path.join(settings.MEDIA_ROOT, "templates")
#     os.makedirs(template_dir, exist_ok=True)

#     template_path = os.path.join(template_dir, template_file.name)
#     with open(template_path, "wb+") as f:
#         for chunk in template_file.chunks():
#             f.write(chunk)

#     ext = os.path.splitext(template_file.name)[1].lower()
#     template_text = extract_template_text(template_path, ext)

#     if not template_text:
#         return JsonResponse({"error": "Unsupported template format"}, status=400)

#     template_words_list = [
#         normalize(w) for w in template_text.split() if normalize(w)
#     ]

#     # -------- STEP 5: SMART MATCHING --------
#     matched_words_raw = []
#     matcher = difflib.SequenceMatcher(
#         None, template_words_list, transcript_words_list
#     )

#     for tag, i1, i2, j1, j2 in matcher.get_opcodes():
#         if tag == "equal":
#             matched_words_raw.extend(transcript[j1:j2])

#     # -------- STEP 6: DB INSERT --------
#     base_name = os.path.splitext(audio_file.name)[0]

#     insert_result = collection.insert_one({
#         "audio": audio_file.name,
#         "template": template_file.name,
#         "language": language,
#         "created_at": datetime.utcnow(),
#         "transcript": transcript,
#         "matched": []
#     })

#     record_id = str(insert_result.inserted_id)

#     # -------- STEP 7: EXPORT WORD CLIPS (ONLY SMALL WORDS KONJM SLOW - CLEAR SOUND) --------
#     clips_dir = os.path.join(settings.MEDIA_ROOT, "audio_clips")
#     os.makedirs(clips_dir, exist_ok=True)

#     original_audio = AudioSegment.from_file(audio_path)

#     # Chinna chinna common words — intha list-la irundha mattum konjam slow pannum (clear-ah kekum)
#     SMALL_COMMON_WORDS = {
#         "and", "or", "with", "the", "a", "an", "in", "on", "at", "to", "for",
#         "of", "but", "so", "if", "is", "as", "by", "from", "that", "this",
#         "it", "not", "be", "are", "was", "were", "have", "has", "had",
#         "you", "i", "we", "they", "he", "she", "me", "us", "them"
#     }

#     for idx, w in enumerate(matched_words_raw):
#         start_ms = int(w["start"] * 1000)
#         end_ms = int(w["end"] * 1000)

#         if end_ms <= start_ms:
#             end_ms = start_ms + 150

#         clip = original_audio[start_ms:end_ms]

#         normalized_word = normalize(w["raw_word"])

#         # Only chinna common words-ah konjam slow pannu (85% speed) → romba clear-ah kekum
#         # Other words full original speed
#         if normalized_word in SMALL_COMMON_WORDS and clip.duration_seconds < 0.4:  # short words only
#             clip = clip._spawn(
#                 clip.raw_data,
#                 overrides={"frame_rate": int(clip.frame_rate * 0.85)}  # konjam slow, but natural
#             ).set_frame_rate(clip.frame_rate)

#         # Safety cap
#         if clip.duration_seconds > 3:
#             clip = clip[:3000]

#         clip_name = f"{base_name}_{record_id}_word_{idx}_{w['word']}.wav"
#         txt_name = clip_name.replace(".wav", ".txt")

#         clip.export(os.path.join(clips_dir, clip_name), format="wav")

#         content = format_time_info(w["start"], w["end"], w["raw_word"])
#         with open(os.path.join(clips_dir, txt_name), "w", encoding="utf-8") as tf:
#             tf.write(content)

#     # -------- STEP 8: GROUPING --------
#     grouped_phrases = []

#     if matched_words_raw:
#         current = [matched_words_raw[0]]

#         for i in range(1, len(matched_words_raw)):
#             prev = current[-1]
#             curr = matched_words_raw[i]

#             if (
#                 curr["index"] == prev["index"] + 1 and
#                 curr["end"] - current[0]["start"] < 20
#             ):
#                 current.append(curr)
#             else:
#                 grouped_phrases.append(current)
#                 current = [curr]

#         grouped_phrases.append(current)

#     # -------- STEP 9: EXPORT PHRASES --------
#     matched_results = []

#     for idx, group in enumerate(grouped_phrases):
#         start = group[0]["start"]
#         end = group[-1]["end"]
#         text = " ".join(w["raw_word"] for w in group)

#         clip = original_audio[int(start * 1000):int(end * 1000)]

#         if clip.duration_seconds > 2:
#             clip = clip[:2000]

#         audio_file_name = f"{base_name}_{record_id}_matched_{idx}.wav"
#         text_file_name = audio_file_name.replace(".wav", ".txt")

#         clip.export(os.path.join(clips_dir, audio_file_name), format="wav")

#         with open(os.path.join(clips_dir, text_file_name), "w", encoding="utf-8") as f:
#             f.write(format_time_info(start, end, text))

#         matched_results.append({
#             "text": text,
#             "start": start,
#             "end": end,
#             "duration": round(end - start, 3),
#             "audio_file": audio_file_name,
#             "text_file": text_file_name
#         })

#     # -------- STEP 10: UPDATE DB --------
#     collection.update_one(
#         {"_id": insert_result.inserted_id},
#         {"$set": {"matched": matched_results}}
#     )

#     # -------- STEP 11: RESPONSE --------
#     return JsonResponse(
#         {
#             "_id": record_id,
#             "transcript": transcript,
#             "matched": matched_results
#         },
#         json_dumps_params={"indent": 2}
#     )

